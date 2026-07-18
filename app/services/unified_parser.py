# -*- coding: utf-8 -*-
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Dict
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client

    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False


@dataclass
class AppPart:
    """Деталь из файла Заявки"""
    name_raw: str
    weight: float
    qty: int


@dataclass
class LayoutPart:
    """Деталь из файла Раскладки - как PartInfo в parser.py"""
    name: str = ""
    dx: float = 0.0
    dy: float = 0.0
    quantity: int = 1


@dataclass
class MergedPart:
    """Итоговая деталь"""
    name: str
    dx: float
    dy: float
    quantity: int
    weight: Optional[float]
    qty_layout: int


@dataclass
class ApplicationData:
    order_name: str = ""
    material: str = "Steel"
    thickness: float = 0.0
    total_weight: Optional[float] = None
    parts: List[AppPart] = field(default_factory=list)
    placed_parts_count: Optional[int] = None
    ordered_parts_count: Optional[int] = None


@dataclass
class LayoutData:
    layout_code: str = "001"
    machine_type: str = "CNF"
    sheet_w: float = 0.0
    sheet_h: float = 0.0
    sheet_weight: Optional[float] = None
    sheet_count: int = 1
    material: str = "Steel"
    thickness: float = 0.0
    cut_time: str = "00:00"
    move_time: str = "00:00"
    pierce_time: str = "00:00"
    total_time: str = "00:00"
    cut_length: Optional[float] = None
    travel_length: Optional[float] = None
    pierces: Optional[int] = None
    cnc_path: str = ""
    parts: List[LayoutPart] = field(default_factory=list)


def _safe_float(val: str) -> Optional[float]:
    try:
        return float(val.replace(',', '.'))
    except Exception:
        return None


def _safe_int(val: str) -> Optional[int]:
    try:
        return int(float(val.replace(',', '.')))
    except Exception:
        return None


def _parse_vertical_parts_block(text: str) -> List[LayoutPart]:
    """
    Парсит «вертикальный» вывод Word для раскладки:
    после заголовка "Имя детали" идут строки:
      <номер>
      <путь D:/.../NAME-4ШТ>
      <DX>
      <DY>
      <Кол-во>
    """
    lines = [ln.strip() for ln in text.split('\n')]
    in_table = False
    current: dict = {}
    parts: List[LayoutPart] = []
    suffix_pattern = re.compile(r'[\s-]*\d+\s*[ШшСс][Тт]\s*$', re.I)

    for raw in lines:
        line = raw.replace('|', '').strip()
        if not line:
            continue

        if 'Имя детали' in line:
            in_table = True
            continue
        if not in_table:
            continue

        if line in ('DX', 'DY', 'Кол-во') or line.startswith('---'):
            continue

        # номер детали — закрываем предыдущую
        if re.fullmatch(r'\d+', line):
            if current.get('name') and current.get('dx') is not None and current.get('dy') is not None:
                parts.append(LayoutPart(
                    name=current['name'],
                    dx=current['dx'],
                    dy=current['dy'],
                    quantity=current.get('qty', 1)
                ))
            current = {}
            continue

        # путь к детали
        if re.match(r'^[A-Za-z]\s*:\\', line):
            full_name = Path(line).name
            clean_name = suffix_pattern.sub('', full_name).strip()
            current['name'] = clean_name
            continue

        # числа (dx, dy, qty)
        val = _safe_float(line)
        if val is None:
            continue
        if 'name' in current:
            if 'dx' not in current:
                current['dx'] = val
            elif 'dy' not in current:
                current['dy'] = val
            elif 'qty' not in current:
                current['qty'] = int(val)

    if current.get('name') and current.get('dx') is not None and current.get('dy') is not None:
        parts.append(LayoutPart(
            name=current['name'],
            dx=current['dx'],
            dy=current['dy'],
            quantity=current.get('qty', 1)
        ))

    return parts


def extract_text(filepath: str) -> str:
    """Извлекает текст из файла Word"""
    try:
        # 1. Windows (локально) - используем win32com
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(Path(filepath).absolute()))
        text = doc.Range().Text
        doc.Close(False)
        word.Quit()
        return clean_text(text)

    except ImportError:
        # 2. Docker (Linux) - используем LibreOffice + python-docx
        import tempfile
        import subprocess
        from docx import Document

        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                # Конвертируем .doc в .docx через LibreOffice
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', tmpdir, filepath],
                    check=True,
                    capture_output=True,
                    timeout=30
                )

                # Ищем созданный файл
                filename = Path(filepath).stem
                docx_path = Path(tmpdir) / f"{filename}.docx"

                if docx_path.exists():
                    doc = Document(docx_path)
                    full_text = []

                    # 1. Читаем обычные параграфы
                    for para in doc.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text.strip())

                    # 2. Читаем таблицы (сохраняем структуру через разделитель |)
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            # Пропускаем пустые строки таблицы
                            if any(cells):
                                full_text.append(" | ".join(cells))

                    return clean_text('\n'.join(full_text))
                else:
                    raise FileNotFoundError("LibreOffice не создал файл .docx")

        except Exception as e:
            # Если LibreOffice не сработал, выводим понятную ошибку
            raise RuntimeError(f"Ошибка извлечения текста через LibreOffice: {e}")


def clean_text(text: str) -> str:
    """Очищает текст от мусора Word, сохраняя структуру строк"""
    # 1. Нормализуем переносы строк
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # 2. Удаляем управляющие символы (кроме \n и \t)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # 3. Удаляем неразрывные пробелы и zero-width space
    text = text.replace('\u00A0', ' ')
    text = text.replace('\u200B', '')
    text = text.replace('\uFEFF', '')

    # 4. Нормализуем пробелы внутри строк
    lines = text.split('\n')
    cleaned = []
    prev_empty = False

    for line in lines:
        # Заменяем множественные пробелы/табы на один пробел
        line = re.sub(r'[ \t]+', ' ', line.strip())

        if line:
            # Непустая строка — сохраняем
            cleaned.append(line)
            prev_empty = False
        elif not prev_empty:
            # Пустая строка, но предыдущая была непустой — сохраняем одну
            cleaned.append('')
            prev_empty = True
            # Иначе пропускаем (не добавляем множественные пустые строки)

    return '\n'.join(cleaned)


def normalize_name(name: str) -> str:
    """Нормализует имя для сопоставления"""
    clean = name.replace('\\', '/').split('/')[-1]
    clean = re.sub(r'\.dft$', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'<[^>]*>', '', clean)
    clean = re.sub(r'^[a-zA-Z]+>\s*', '', clean)
    clean = re.sub(r'\s+', ' ', clean).strip()
    clean = clean.lower()
    clean = clean.replace('мм', 'mm').replace('шт', 'sht')
    return clean


def parse_application_text(text: str) -> ApplicationData:
    """Парсит файл Заявки (.DOC)"""
    data = ApplicationData()

    # Заказ: + следующая строка (Word-вывод) ИЛИ таблица с |
    order_match = re.search(r'Заказ\s*:\s*\n\s*([^\s|]+)', text)
    if order_match:
        data.order_name = order_match.group(1).strip()
        data.order_name = re.sub(r'\.dsp$', '', data.order_name, flags=re.I)

    mat_match = re.search(r'Материал\s*:\s*([^\s|]+)', text)
    if mat_match:
        data.material = mat_match.group(1).strip()

    # Толщина: пытаемся взять из вертикального блока "Субраскладки в заказе"
    thickness_match = re.search(r'Субраскладки в заказе[\s\S]*?\nSteel\s*\n(\d+(?:[.,]\d+)?)\s*\n', text, re.I)
    if thickness_match:
        th = _safe_float(thickness_match.group(1))
        if th is not None:
            data.thickness = th

    # Толщина из pipe-таблицы: ищем заголовок с "Толщ", берём индекс столбца
    # Важно: split('|') БЕЗ фильтрации пустых, чтобы сохранить выравнивание
    if data.thickness == 0.0:
        lines_tmp = text.split('\n')
        for li, ln in enumerate(lines_tmp):
            if 'Толщ' in ln and ('Материал' in ln or 'Material' in ln):
                # Индекс столбца "Толщ" в заголовке (считая пустые ячейки)
                header_parts = [h.strip() for h in ln.split('|')]
                th_col = None
                for hi, h in enumerate(header_parts):
                    if 'Толщ' in h:
                        th_col = hi
                        break
                # Следующая строка — данные
                if th_col is not None and li + 1 < len(lines_tmp):
                    data_line = lines_tmp[li + 1]
                    data_parts = [c.strip() for c in data_line.split('|')]
                    if th_col < len(data_parts):
                        try:
                            th = _safe_float(data_parts[th_col])
                            if th is not None and th > 0:
                                data.thickness = th
                                break
                        except (ValueError, IndexError):
                            pass

    # ========== Толщина из таблицы с | (LibreOffice/docx) ==========
    lines = text.split('\n')

    # Ищем таблицу "Данные материала" или "Доп. данные материала"
    in_material_table = False
    header_found = False

    for i, line in enumerate(lines):
        if 'Данные материала' in line or 'Доп. данные материала' in line:
            in_material_table = True
            continue

        if in_material_table and line.startswith('| ---'):
            header_found = True
            continue

        if in_material_table and header_found:
            if line.startswith('|'):
                cells = [c.strip() for c in line.split('|') if c.strip()]
                # В "Данные материала": | Steel | 8 | ... |
                # Или: | Stainless | 50 | ... |
                if len(cells) >= 2:
                    try:
                        thickness_val = cells[1].replace(',', '.').strip()
                        if thickness_val and thickness_val not in ['.', ','] and thickness_val.replace('.', '').isdigit():
                            data.thickness = float(thickness_val)
                    except (ValueError, IndexError):
                        pass
                if data.thickness > 0:
                    break

            # Если дошли до следующей таблицы — выходим
            if 'Детали в субраскладках' in line:
                break

    # ========== ИСПРАВЛЕНИЕ 2: Общий вес ==========
    # Ищем "Общий в ес ( KG )" в таблице
    for i, line in enumerate(lines):
        if 'Общий' in line and 'вес' in line and 'KG' in line:
            # Строка-заголовок содержит "Общий вес (KG)" — находим индекс столбца
            header_parts = [c.strip() for c in line.split('|')]
            weight_col = None
            for ci, cell in enumerate(header_parts):
                if 'Общий' in cell and 'вес' in cell:
                    weight_col = ci
                    break

            # Проверяем текущую строку (если значение в той же строке)
            for cell in [c.strip() for c in line.split('|')]:
                try:
                    val = float(cell.replace(',', '.'))
                    if 10 < val < 100000:
                        data.total_weight = val
                        break
                except ValueError:
                    pass

            # Если не нашли — смотрим следующую строку (pipe-таблица без | в начале)
            if data.total_weight is None and i + 1 < len(lines):
                next_line = lines[i + 1]
                data_parts = [c.strip() for c in next_line.split('|')]
                if weight_col is not None and weight_col < len(data_parts):
                    try:
                        val = float(data_parts[weight_col].replace(',', '.'))
                        if 10 < val < 100000:
                            data.total_weight = val
                    except (ValueError, IndexError):
                        pass
                elif data_parts:
                    try:
                        val = float(data_parts[0].replace(',', '.'))
                        if 10 < val < 100000:
                            data.total_weight = val
                    except ValueError:
                        pass
            break

    # ========== Парсинг деталей ==========
    in_parts_table = False
    parts_vertical: List[AppPart] = []
    current_name: Optional[str] = None
    current_weight: Optional[float] = None
    current_qty: Optional[int] = None

    for line in lines:
        if 'Детали в субраскладках' in line:
            in_parts_table = True
            continue

        if not in_parts_table:
            continue

        # заголовки / разделители
        if line.startswith('| ---') or line.startswith('| №'):
            continue
        if line in ('№', 'Имя файла детали', 'Вид', 'Материал', 'Толщ. (мм)', 'Вес (KG', 'Вес (KG)', 'Кол-во', 'Кол-во ', 'Кол-во\n'):
            continue

        cells = [c.strip() for c in line.split('|') if c.strip()]

        # 1) Формат LibreOffice/docx (таблица с |)
        if len(cells) >= 7:
            try:
                name_raw = cells[1]
                if '.dft' in name_raw.lower() or 'шт' in name_raw.lower():
                    # Пустая ячейка "Вид" фильтруется → сдвиг индексов.
                    # 8 ячеек (с пустой "Вид"): weight=5, qty=6
                    # 7 ячеек (пустая "Вид" отфильтрована): weight=4, qty=5
                    if len(cells) == 7:
                        weight_str = cells[4]
                        qty_str = cells[5]
                    else:
                        weight_str = cells[5]
                        qty_str = cells[6]

                    data.parts.append(AppPart(
                        name_raw=name_raw,
                        weight=float(weight_str.replace(',', '.')),
                        qty=int(qty_str)
                    ))
            except (ValueError, IndexError):
                continue
            continue

        # 2) Вертикальный формат Word: имя детали (строка с путём/именем) + далее вес и кол-во
        # Пример имён в заявке: 8мм-1-616-2шт (без путей) — оставляем как есть.
        if re.search(r'[А-ЯA-Z0-9].*-\d+\s*[ШшСс][Тт]$', line):
            # закрываем предыдущую, если собрана
            if current_name and current_weight is not None and current_qty is not None:
                parts_vertical.append(AppPart(name_raw=current_name, weight=current_weight, qty=current_qty))
            current_name = line.strip()
            current_weight = None
            current_qty = None
            continue

        # иногда имя может быть с путём
        if re.match(r'^[A-Za-z]\s*:\\', line):
            if current_name and current_weight is not None and current_qty is not None:
                parts_vertical.append(AppPart(name_raw=current_name, weight=current_weight, qty=current_qty))
            current_name = Path(line).name
            current_weight = None
            current_qty = None
            continue

        # Вес (KG) и Кол-во — обычно числа в отдельных строках после имени
        num = _safe_float(line)
        if num is None:
            continue
        if current_name:
            # по заявке вес обычно дробный, кол-во — целое
            if current_weight is None:
                current_weight = num
            elif current_qty is None:
                current_qty = int(num)

    if current_name and current_weight is not None and current_qty is not None:
        parts_vertical.append(AppPart(name_raw=current_name, weight=current_weight, qty=current_qty))

    if not data.parts and parts_vertical:
        data.parts = parts_vertical

    # ========== Размещено / Заказано деталей ==========
    placed_match = re.search(r'Размещено\s+деталей.*?\|\s*(\d+)', text, re.I)
    if not placed_match:
        placed_match = re.search(r'Размещено\s+деталей.*?:\s*(\d+)', text, re.I)
    if placed_match:
        data.placed_parts_count = int(placed_match.group(1))

    ordered_match = re.search(r'Заказано\s+деталей.*?\|\s*(\d+)', text, re.I)
    if not ordered_match:
        ordered_match = re.search(r'Заказано\s+деталей.*?:\s*(\d+)', text, re.I)
    if ordered_match:
        data.ordered_parts_count = int(ordered_match.group(1))

    return data


def parse_layout_text(text: str, filename: str) -> LayoutData:
    """Парсит файл Раскладки (.Cnf/.Fnf.DOC)"""
    # Очищаем текст
    text = clean_text(text)

    print(f"\n🔍 DEBUG parse_layout_text:")
    print(f"   Получено {len(text)} символов")
    print(f"   Первые 1000 символов:\n{text[:1000]}\n")

    data = LayoutData()

    if "fnf" in filename.lower():
        data.machine_type = "FNF"

    code_match = re.search(r'(\d{3})', filename)
    if code_match:
        data.layout_code = code_match.group(1)

    # Базовые поля из раскладки (по xlsx)
    sheet_count_match = re.search(r'Количество\s+листов\s*:\s*(\d+)', text, re.I)
    if sheet_count_match:
        sc = _safe_int(sheet_count_match.group(1))
        if sc is not None and sc > 0:
            data.sheet_count = sc

    mat_match = re.search(r'Материал\s*:\s*([^\n]+)', text, re.I)
    if mat_match:
        data.material = mat_match.group(1).strip()

    thick_match = re.search(r'Толщина\s*:\s*([\d.,]+)', text, re.I)
    if thick_match:
        th = _safe_float(thick_match.group(1))
        if th is not None:
            data.thickness = th

    weight_match = re.search(r'Вес\s*:\s*([\d.,]+)', text, re.I)
    if weight_match:
        w = _safe_float(weight_match.group(1))
        if w is not None:
            data.sheet_weight = w

    total_time_match = re.search(r'Время,\s*всего\s*:\s*([0-9]{1,3}:[0-9]{2})', text, re.I)
    if total_time_match:
        data.total_time = total_time_match.group(1)

    travel_match = re.search(r'Перемещ\.\s*\(мм\)\s*:\s*([\d.,]+)', text, re.I)
    if travel_match:
        data.travel_length = _safe_float(travel_match.group(1))

    cut_len_match = re.search(r'Резка\s*\(мм\)\s*:\s*([\d.,]+)', text, re.I)
    if cut_len_match:
        data.cut_length = _safe_float(cut_len_match.group(1))

    pierces_match = re.search(r'Кол\.?\s*проколов\s*:\s*(\d+)', text, re.I)
    if pierces_match:
        data.pierces = _safe_int(pierces_match.group(1))

    size_match = re.search(r'Размер\s*:\s*([\d.,]+)\s*[XxхХ]\s*([\d.,]+)', text, re.I)
    if size_match:
        data.sheet_w = float(size_match.group(1).replace(',', '.'))
        data.sheet_h = float(size_match.group(2).replace(',', '.'))
        print(f"   ✅ Размер листа: {data.sheet_w}x{data.sheet_h}")

    for pattern, attr in [
        (r'Перемещение\s*:\s*([\d:]+)', 'move_time'),
        (r'Резка\s*:\s*([\d:]+)', 'cut_time'),
        (r'Прокалывание\s*:\s*([\d:]+)', 'pierce_time')
    ]:
        match = re.search(pattern, text)
        if match:
            setattr(data, attr, match.group(1))

    cnc_match = re.search(r'Имя файла УП\s*:\s*(\S+)', text)
    if cnc_match:
        data.cnc_path = cnc_match.group(1).strip()

    # ========== ПАРСИНГ ДЕТАЛЕЙ ==========
    # 1) Попытка через «табличные» строки с |
    parts = []
    in_table = False
    cur = {}  # Текущая собираемая деталь

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue

        if 'Имя детали' in line:
            in_table = True
            print(f"   ✅ Найдена таблица деталей")
            continue

        if not in_table:
            continue

        # Разбиваем строку по разделителю |
        # Это даст нам доступ к отдельным колонкам таблицы
        cells = [c.strip() for c in line.split('|') if c.strip()]

        # Если ячеек меньше 2, скорее всего это мусор или заголовок
        if len(cells) < 2:
            continue

        # 1. Ищем ячейку, которая содержит путь (D:\ или C:\)
        path_cell = None
        for cell in cells:
            if re.search(r'[A-Za-z]:\\', cell):
                path_cell = cell
                break

        if path_cell:
            # Если нашли путь — это начало новой детали

            # Сначала сохраняем предыдущую деталь, если она была
            if cur.get('name') and cur.get('dx') is not None:
                parts.append(LayoutPart(
                    name=cur['name'],
                    dx=cur['dx'],
                    dy=cur['dy'],
                    quantity=cur.get('qty', 1)
                ))

            # Извлекаем чистое имя файла ИЗ ЯЧЕЙКИ с путем
            full_name = path_cell.replace('/', '\\').split('\\')[-1]
            clean_name = re.sub(r'\.dft$', '', full_name, flags=re.I).strip()

            # Инициализируем новую деталь
            cur = {
                'name': clean_name,
                'dx': None,
                'dy': None,
                'qty': None
            }

            print(f"   📄 Обнаружена деталь: {clean_name}")

            # 2. Пытаемся найти числа (DX, DY, QTY) в ячейках СПРАВА от пути
            # Находим индекс ячейки с путем
            try:
                path_idx = cells.index(path_cell)
                # Смотрим ячейки после пути
                numbers = []
                for i in range(path_idx + 1, len(cells)):
                    try:
                        val = float(cells[i].replace(',', '.'))
                        numbers.append(val)
                    except ValueError:
                        pass

                # Если нашли 3 числа подряд — записываем их
                if len(numbers) >= 3:
                    cur['dx'] = numbers[0]
                    cur['dy'] = numbers[1]
                    cur['qty'] = int(numbers[2])
                    print(f"   ✅ Сразу найдены размеры: {cur['dx']}x{cur['dy']} кол-во {cur['qty']}")
                elif len(numbers) == 2:
                    cur['dx'] = numbers[0]
                    cur['dy'] = numbers[1]
                elif len(numbers) == 1:
                    cur['dx'] = numbers[0]

            except ValueError:
                pass  # Если не нашли индекс, ничего страшного

            continue

        # 3. Если пути в строке нет, но у нас есть "незавершенная" деталь
        if cur.get('name') and cur.get('dx') is None:
            # Пытаемся найти числа в любой ячейке этой строки
            for cell in cells:
                try:
                    val = float(cell.replace(',', '.'))
                    if cur.get('dx') is None:
                        cur['dx'] = val
                    elif cur.get('dy') is None:
                        cur['dy'] = val
                    elif cur.get('qty') is None:
                        cur['qty'] = int(val)
                except ValueError:
                    pass

    # Сохраняем последнюю деталь
    if cur.get('name') and cur.get('dx') is not None:
        parts.append(LayoutPart(
            name=cur['name'],
            dx=cur['dx'],
            dy=cur['dy'],
            quantity=cur.get('qty', 1)
        ))

    data.parts = parts
    # 2) Если в Word-выводе не было | и ничего не нашли — fallback на вертикальный парсинг
    if not data.parts:
        data.parts = _parse_vertical_parts_block(text)

    print(f"\n   ✅ ВСЕГО: Найдено {len(data.parts)} деталей")
    for idx, p in enumerate(data.parts, 1):
        print(f"      {idx}. {p.name} | DX:{p.dx} DY:{p.dy} QTY:{p.quantity}")

    return data


def extract_images(filepath: str, output_dir: str, prefix: str = "", filter_dft: bool = False) -> List:
    """Извлекает изображения из DOC файла.
    filter_dft=True — только изображения рядом с .dft именами (для заявок).
      Возвращает список кортежей (image_path, dft_name) для маппинга по имени.
    filter_dft=False — все изображения (для раскладок).
      Возвращает список путей [image_path, ...].
    """
    import tempfile
    import subprocess
    from pathlib import Path
    import shutil

    saved = []
    IMG_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.emf', '.wmf', '.tiff', '.tif'}

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'html', '--outdir', tmpdir, filepath],
                check=True,
                capture_output=True,
                timeout=60
            )

            dest_dir = Path(output_dir) / prefix
            dest_dir.mkdir(parents=True, exist_ok=True)

            html_files = list(Path(tmpdir).glob('*.html'))
            if not html_files:
                return saved

            html_text = html_files[0].read_text(encoding='utf-8', errors='ignore')

            from urllib.parse import unquote

            # Собираем ВСЕ изображения из tmpdir
            all_images = {}
            html_name = Path(filepath).stem

            # 1. Ищем в ВСЕХ папках *_files (LibreOffice может создавать разные имена)
            for img_dir in Path(tmpdir).glob("*_files"):
                if img_dir.is_dir():
                    for img_file in img_dir.iterdir():
                        if img_file.is_file() and img_file.suffix.lower() in IMG_EXTS:
                            all_images[img_file.name] = img_file
                            print(f"DEBUG extract_images: found in _files: {img_file.name}", flush=True)

            # 2. Ищем в корне tmpdir (LibreOffice может класть GIF рядом с HTML)
            for img_file in Path(tmpdir).iterdir():
                print(f"DEBUG extract_images: checking {img_file.name}, is_file={img_file.is_file()}, ext={img_file.suffix.lower()}", flush=True)
                if img_file.is_file() and img_file.suffix.lower() in IMG_EXTS:
                    if img_file.name not in all_images:
                        all_images[img_file.name] = img_file
                        print(f"DEBUG extract_images: found in root: {img_file.name}", flush=True)
            print(f"DEBUG extract_images: all_images keys = {list(all_images.keys())}", flush=True)

            # Для каждого <img> в HTML
            for m in re.finditer(r'<img[^>]+src=["\']([^"\']+)', html_text):
                img_ref = unquote(m.group(1)).split('/')[-1]
                print(f"DEBUG extract_images: img_ref={img_ref}, looking in all_images", flush=True)
                src = all_images.get(img_ref)
                if not src:
                    print(f"DEBUG extract_images: NOT FOUND: {img_ref}", flush=True)
                    continue
                print(f"DEBUG extract_images: FOUND: {img_ref} -> {src}", flush=True)

                if filter_dft:
                    img_pos = m.start()
                    context = html_text[max(0, img_pos - 3000):img_pos + 200]
                    stripped = re.sub(r'<[^>]*>', '', context)
                    dft_pos = stripped.lower().rfind('.dft')
                    if dft_pos <= 0:
                        continue
                    before = stripped[max(0, dft_pos - 200):dft_pos]
                    lines = before.split('\n')
                    name_lines = []
                    for line in reversed(lines):
                        line = line.strip()
                        if not line or line.startswith('|') or line.startswith('---'):
                            break
                        name_lines.insert(0, line)
                    name = ' '.join(name_lines).strip()
                    if not name:
                        continue
                    dest_path = dest_dir / img_ref
                    shutil.copy2(src, dest_path)
                    saved.append((f"/api/v1/images/{prefix}/{img_ref}", name))
                else:
                    dest_path = dest_dir / img_ref
                    shutil.copy2(src, dest_path)
                    saved.append(f"/api/v1/images/{prefix}/{img_ref}")

    except Exception as e:
        print(f"Warning: image extraction failed: {e}")

    return saved


def extract_layout_image(filepath: str, output_dir: str, prefix: str = "", dpi: int = 300) -> Optional[str]:
    """Извлекает изображение раскладки через LibreOffice (fallback — теряет кривые).

    Returns: URL изображения или None при ошибке.
    """
    images = extract_images(filepath, output_dir, prefix=prefix, filter_dft=False)
    if images:
        return images[0]
    return None


def merge_data(app_data: ApplicationData, layout_data: LayoutData) -> List[MergedPart]:
    """Объединяет данные из Заявки и Раскладки"""
    if not app_data.parts:
        return [
            MergedPart(
                name=p.name,
                dx=p.dx,
                dy=p.dy,
                quantity=p.quantity,
                weight=None,
                qty_layout=p.quantity
            )
            for p in layout_data.parts
        ]

    # Создаем словарь из заявки для быстрого поиска
    app_dict: Dict[str, AppPart] = {}
    for p in app_data.parts:
        key = normalize_name(p.name_raw)
        app_dict[key] = p

    result = []
    for lp in layout_data.parts:
        key = normalize_name(lp.name)
        ap = app_dict.get(key)

        if ap:
            # Есть совпадение
            result.append(MergedPart(
                name=ap.name_raw.replace('.dft', '').replace('.DFT', ''),
                dx=lp.dx,
                dy=lp.dy,
                quantity=ap.qty,
                weight=ap.weight,
                qty_layout=lp.quantity
            ))
        else:
            # Нет совпадения
            result.append(MergedPart(
                name=lp.name,
                dx=lp.dx,
                dy=lp.dy,
                quantity=lp.quantity,
                weight=None,
                qty_layout=lp.quantity
            ))

    return result


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Использование: python unified_parser.py <путь_к_файлу.doc>")
        sys.exit(1)

    filepath = sys.argv[1]
    filename = Path(filepath).name

    print(f"📂 Парсинг: {filepath}")

    try:
        # Извлекаем текст
        text = extract_text(filepath)
        print("✅ Текст извлечён")

        # Парсим раскладку
        layout_data = parse_layout_text(text, filename)
        print(f"✅ Найдено деталей: {len(layout_data.parts)}")

        # Выводим информацию
        print(f"\n📊 Основная информация:")
        print(f"   Тип: {layout_data.machine_type}")
        print(f"   Код: {layout_data.layout_code}")
        print(f"   Размер листа: {layout_data.sheet_w} x {layout_data.sheet_h}")
        print(f"   CNC путь: {layout_data.cnc_path}")
        print(f"   Время резки: {layout_data.cut_time}")
        print(f"   Время перемещения: {layout_data.move_time}")

        print(f"\n🔧 Детали:")
        for i, part in enumerate(layout_data.parts, 1):
            print(f"   {i}. {part.name} | DX:{part.dx} DY:{part.dy} QTY:{part.quantity}")

    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback

        traceback.print_exc()
        sys.exit(1)